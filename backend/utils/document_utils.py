from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import base64
from datetime import datetime

def create_chat_document(messages):
    """Create a Word document from chat messages."""
    doc = Document()
    
    # Add title
    title = doc.add_heading('Chat Conversation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add timestamp
    timestamp = doc.add_paragraph()
    timestamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    timestamp_run = timestamp.add_run(f'Generated on {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    timestamp_run.font.size = Pt(10)
    timestamp_run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Add a line break
    doc.add_paragraph()
    
    # Add messages
    for msg in messages:
        # Add sender label
        sender = "AI Assistant" if msg['sender'] == 'ai' else "User"
        p = doc.add_paragraph()
        sender_run = p.add_run(f"{sender}:")
        sender_run.bold = True
        if msg['sender'] == 'ai':
            sender_run.font.color.rgb = RGBColor(0, 112, 192)  # Blue for AI
        else:
            sender_run.font.color.rgb = RGBColor(46, 116, 46)  # Green for User
            
        # Add message text
        doc.add_paragraph(msg['text'])
        
        # Add spacing between messages
        doc.add_paragraph()
    
    # Save to bytes
    doc_bytes = bytes()
    from io import BytesIO
    doc_stream = BytesIO()
    doc.save(doc_stream)
    doc_bytes = doc_stream.getvalue()
    
    # Convert to base64
    doc_base64 = base64.b64encode(doc_bytes).decode('utf-8')
    
    return doc_base64 